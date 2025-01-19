from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches
from io import BytesIO
from PIL import Image
import fitz
from datetime import datetime
import tkinter as tk

# root = tk.Tk()

def get_cropped_Image(original_image, page_number):
    l = [1, 5, 6, 7, 8, 9, 12, 23, 25, 26, 27, 28, 31, 42]
    l1 = [3, 14]
    if page_number in l:
        return original_image.crop((0, 20, 612, 567))
    elif page_number in l1:
        return original_image.crop((0, 20, 612, 637))
    else:
        return original_image.crop((0, 20, 612, 772))


def convert_pdf_pages_to_images(pdf_path):
    images = []
    try:
        doc = fitz.open(pdf_path)
        for page_num in range(1, doc.page_count):
            if page_num < doc.page_count:
                page = doc.load_page(page_num)
                pix = page.get_pixmap()
                img_data = pix.tobytes("png")
                img = Image.open(BytesIO(img_data))
                cropped_img = get_cropped_Image(img, page_num)

                img_stream = BytesIO()
                cropped_img.save(img_stream, format="PNG")
                img_stream.seek(0)
                images.append((img_stream, cropped_img.size))  # Save the stream and original image size
        doc.close()
    except Exception as e:
        print(f"An error occurred while converting PDF to images: {e}")
    return images


def add_header_and_footer(doc):
    section = doc.sections[-1]
    
    # Header
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = "Created by Pratham Kubetkar"
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = header_paragraph.runs[0]
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    # Footer
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = f"Created by Pratham Kubetkar on {datetime.now().strftime('%d-%m-%Y %H:%M:%S')}"
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = footer_paragraph.runs[0]
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)


def create_custom_page(doc):
    section = doc.sections[-1]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture("C:/Users/Pratham/Pictures/4217614.jpg", width=Inches(8.5), height=Inches(11))
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("\n\nCustom Designed Page", style="Title").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    content = doc.add_paragraph()
    run = content.add_run("This is a customizable page template. You can add more designs or information here.")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)


def create_docx_from_images(images, docx_output_path):
    try:
        doc = Document()
        add_header_and_footer(doc)  # Add header and footer to the document
        max_width = Inches(6.5)

        create_custom_page(doc)

        for img_stream, img_size in images:
            width, height = img_size
            aspect_ratio = height / width

            target_width = max_width
            target_height = target_width * aspect_ratio

            doc.add_picture(img_stream, width=target_width, height=target_height)

        doc.save(docx_output_path)
        print(f"DOCX file created successfully: {docx_output_path}")
    except Exception as e:
        print(f"An error occurred while creating the DOCX file: {e}")


if __name__ == "__main__":
    pdf_path = "C:/Users/Pratham/Downloads/Kundali2.pdf"
    docx_output_path = "C:/Users/Pratham/Desktop/MP/EV/Project/output.docx"

    images = convert_pdf_pages_to_images(pdf_path)

    if images:
        create_docx_from_images(images, docx_output_path)
    else:
        print("No images were extracted from the specified PDF pages.")
